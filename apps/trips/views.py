from django.views.generic import ListView, DetailView, CreateView, UpdateView
from django.contrib.auth.mixins import LoginRequiredMixin
from django.urls import reverse_lazy
from django.shortcuts import redirect
from .models import Trip
from .forms import TripForm, TripStatusForm
from apps.users.models import User
from django.conf import settings
from django.db.models import Q
from apps.core.models import Status

CLOSED_STATUSES = [
    'Выполнено',
    'Закрыто',
    'Отменено',
    'Completed',
    'Closed',
]

class TripListView(LoginRequiredMixin, ListView):
    model = Trip
    template_name = 'trips/list.html'
    context_object_name = 'trips'
    ordering = ['-created_at']
    paginate_by = 15

    def get_queryset(self):
        queryset = super().get_queryset()

        # 🔒 Filtering for regular users
        if self.request.user.role == User.Role.USER:
            queryset = queryset.filter(responsible=self.request.user).exclude(status__name__in=CLOSED_STATUSES)

        q = self.request.GET.get('q')
        status_id = self.request.GET.get('status')
        req_num = self.request.GET.get('req_num')

        if q:
            queryset = queryset.filter(
                Q(request_number__icontains=q) |
                Q(description__icontains=q) |
                Q(responsible__full_name__icontains=q) |
                Q(escort_name__icontains=q)
            )

        if req_num:
            if req_num.isdigit():
                padded_num = req_num.zfill(4)
                queryset = queryset.filter(request_number__icontains=padded_num)
            else:
                queryset = queryset.filter(request_number__icontains=req_num)
        
        if status_id:
            queryset = queryset.filter(status_id=status_id)
            
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['statuses'] = Status.objects.all()
        return context

class TripDetailView(LoginRequiredMixin, DetailView):
    model = Trip
    template_name = 'trips/detail.html'
    context_object_name = 'trip'

class TripCreateView(LoginRequiredMixin, CreateView):
    model = Trip
    form_class = TripForm
    template_name = 'trips/form.html'
    success_url = reverse_lazy('trips:list')

    def dispatch(self, request, *args, **kwargs):
        if request.user.role == User.Role.USER:
            return redirect('trips:list')
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['institution_mapping'] = list(Institution.objects.values('id', 'region_id'))
        return context

class TripUpdateView(LoginRequiredMixin, UpdateView):
    model = Trip
    template_name = 'trips/form.html'
    success_url = reverse_lazy('trips:list')

    def dispatch(self, request, *args, **kwargs):
        obj = self.get_object()
        if request.user.role == User.Role.USER:
            closed_statuses = ['Выполнено', 'Закрыто', 'Отменено', 'Completed', 'Closed']
            if obj.status.name in closed_statuses:
                return redirect('trips:list')
        return super().dispatch(request, *args, **kwargs)

    def get_form_class(self):
        if self.request.user.role == User.Role.USER:
            return TripStatusForm
        return TripForm

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['institution_mapping'] = list(Institution.objects.values('id', 'region_id'))
        # Carry over original context if it exists (Status is used in trip detail/update?)
        # Base class ListView for statuses is in TripListView, not update.
        # But TripListView has context['statuses']
        return context


import openpyxl
from django.http import HttpResponse
from django.shortcuts import render, get_object_or_404
from django.core.exceptions import PermissionDenied
from apps.core.models import Region, Institution, EquipmentType, Status
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from django.utils.translation import gettext as _

def export_trips_to_excel(request):
    """
    Export all trips to an Excel file.
    """
    if request.user.role == User.Role.USER:
        raise PermissionDenied

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = 'attachment; filename=trips_export.xlsx'

    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = _('Trips')

    # Header row
    columns = [
        _('Trip Number'), _('Created At'), _('Region'), _('Institution'),
        _('Equipment Type'), _('Status'), _('Responsible'), _('Contact Phone'), 
        _('Description'), _('Escort Name'), _('Escort Phone'), _('Order Number'), _('Result')
    ]
    worksheet.append(columns)

    for trip in Trip.objects.all().order_by('-created_at'):
        row = [
            trip.request_number,
            trip.created_at,
            trip.region.name if trip.region else '',
            trip.institution.name if trip.institution else '',
            trip.equipment_type.name if trip.equipment_type else '',
            trip.status.name if trip.status else '',
            trip.responsible.full_name if trip.responsible else '',
            trip.contact_phone,
            trip.description,
            trip.escort_name,
            trip.escort_phone,
            trip.order_number,
            trip.result_info
        ]
        worksheet.append(row)

    workbook.save(response)
    return response


def import_trips_from_excel(request):
    """
    Import trips from an Excel file.
    """
    if request.user.role == User.Role.USER:
        raise PermissionDenied

    if request.method == 'POST' and request.FILES.get('file'):
        excel_file = request.FILES['file']
        wb = openpyxl.load_workbook(excel_file)
        worksheet = wb.active

        # Start from second row
        for row in worksheet.iter_rows(min_row=2, values_only=True):
            # Maps to export:
            # 0: Number, 1: Created, 2: Region, 3: Institution, 4: Equip, 
            # 5: Status, 6: Resp, 7: Phone, 8: Desc, 9: EscortName, 
            # 10: EscortPhone, 11: Order, 12: Result
            
            region_name = row[2]
            institution_name = row[3]
            equipment_name = row[4]
            status_name = row[5]
            phone = row[7]
            desc = row[8]
            escort_name = row[9]
            escort_phone = row[10]
            order_num = row[11]
            result = row[12]

            region = Region.objects.filter(name__iexact=region_name).first()
            institution = Institution.objects.filter(name__iexact=institution_name).first()
            equipment = EquipmentType.objects.filter(name__iexact=equipment_name).first()
            status = Status.objects.filter(name__iexact=status_name).first()

            if region and institution and equipment and status:
                Trip.objects.create(
                    region=region,
                    institution=institution,
                    equipment_type=equipment,
                    status=status,
                    contact_phone=phone or '',
                    description=desc or '',
                    responsible=request.user,
                    escort_name=escort_name or '',
                    escort_phone=escort_phone or '',
                    order_number=order_num or '',
                    result_info=result or ''
                )
        
        return redirect('trips:list')

    return render(request, 'common/import_form.html')


# def download_trip_docx(request, pk):
#     """
#     Download a single trip as a DOCX file.
#     """
#     trip = get_object_or_404(Trip, pk=pk)
#     with open(settings.BASE_DIR / 'templates/docs/trip_template.docx', 'rb') as f:
#         document = Document(f)
#     replacements = {
#         "{{ESCORT_NAME}}": trip.escort_name,
#         "{{INSTITUTION_NAME}}": trip.institution.name,
#         "{{PASSPORT}}": ""
#     }
#     for paragraph in document.paragraphs:
#         for key, value in replacements.items():
#             if key in paragraph.text:
#                 paragraph.text = paragraph.text.replace(key, value)
#     document.save(response)
#     response = HttpResponse(
#         content_type=(
#             "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )
#     )
#     response["Content-Disposition"] = (
#         f'attachment; filename="trip_{trip.id}.docx"'
#     )
#     return response
def replace_in_paragraph(paragraph, replacements):
    full_text = ''.join(run.text for run in paragraph.runs)

    for key, value in replacements.items():
        if key in full_text:
            full_text = full_text.replace(key, value)

            # clear existing runs
            for run in paragraph.runs:
                run.text = ""

            # put text into first run (keeps its style)
            paragraph.runs[0].text = full_text

def download_trip_docx(request, pk):
    """
    Download a single trip as a DOCX file.
    """
    trip = get_object_or_404(Trip, pk=pk)

    template_path = settings.BASE_DIR / 'templates/docs/trip_template.docx'
    document = Document(template_path)

    replacements = {
        "{{ESCORT_NAME}}": trip.responsible.full_name or "",
        "{{INSTITUTION_NAME}}": trip.institution.name if trip.institution else "",
        "{{PASSPORT}}": trip.responsible.passport or "",
    }
    
    for p in document.paragraphs:
        replace_in_paragraph(p, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, replacements)
    response = HttpResponse(
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    )
    response["Content-Disposition"] = (
        f'attachment; filename="trip_{trip.id}.docx"'
    )

    document.save(response)
    return response