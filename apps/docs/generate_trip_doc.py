from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# создаём документ
doc = Document()

# настраиваем поля страницы
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(1.5)

paragraph = doc.add_paragraph("КОМАНДИРОВОЧНОЕ УДОСТОВЕРЕНИЕ",)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph.runs[0].bold = True
paragraph.runs[0].font.size = Pt(10)
paragraph.runs[0].font.name = "Tahoma"

table = doc.add_table(rows=6, cols=2)
table.autofit = False
cell_0_0 = table.cell(0, 0)
paragraph_0_0 = cell_0_0.text = "Выдано"
# paragraph_0_0.alignment = WD_ALIGN_PARAGRAPH.CENTER
# paragraph_0_0.runs[0].bold = True
# paragraph_0_0.runs[0].font.size = Pt(10)
# paragraph_0_0.runs[0].font.name = "Tahoma"
tc_0_0 = cell_0_0._tc
tcPr_0_0 = tc_0_0.get_or_add_tcPr()

tcW_0_0 = OxmlElement('w:tcW')
tcW_0_0.set(qn('w:w'), str(int(Inches(0.91).pt * 20)))  # twips
tcW_0_0.set(qn('w:type'), 'dxa')

tcPr_0_0.append(tcW_0_0)


cell_0_1 = table.cell(0, 1)
cell_0_1.text = "AXMADJANOV OZODBEK ALISHERJON O’G’LI– инженер по сборке оборудованием"
cell_0_1.alignment = WD_ALIGN_PARAGRAPH.CENTER

tc_0_1 = cell_0_1._tc
tcPr_0_1 = tc_0_1.get_or_add_tcPr()

tcW_0_1 = OxmlElement('w:tcW')
tcW_0_1.set(qn('w:w'), str(int(Inches(6.21).pt * 20)))  # twips
tcW_0_1.set(qn('w:type'), 'dxa')

tcPr_0_1.append(tcW_0_1)

cell_1_0 = table.cell(1, 0)
cell_1_0.text = ""
cell_1_0.width = Inches(0.91)

cell_1_1 = table.cell(1, 1)
cell_1_1.text = "(фамилия, имя, отчество/должность)"
cell_1_1.width = Inches(6.21)
cell_1_1.alignment = WD_ALIGN_PARAGRAPH.CENTER

# сохраняем файл
doc.save("document_with_margins_and_table.docx")
